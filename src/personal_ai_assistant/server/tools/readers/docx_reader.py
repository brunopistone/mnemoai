"""DOCX reading functionality."""

from .. import validate_file_path
import json
from docx import Document
import os

from personal_ai_assistant.utils.config import config
from personal_ai_assistant.utils.logger import logger
from .chunking_helper import __count_tokens as count_tokens, process_large_content

try:
    from ..rag.session import get_rag_session

    _rag_available = True
except ImportError:
    _rag_available = False
    logger.debug("RAG module not available")


async def read_docx(file_path: str) -> str:
    """Read DOCX content with automatic chunking and summarization for large files.

    Args:
        file_path: Path to the DOCX file (use EXACTLY as provided - do NOT escape spaces with backslashes)

    Returns:
        JSON with DOCX content
    """
    try:
        # Validate and normalize path
        is_valid, normalized_path, error_dict = validate_file_path(file_path)
        if not is_valid:
            return json.dumps(error_dict)

        # Check if it's a DOCX file
        if not normalized_path.lower().endswith(".docx"):
            return json.dumps(
                {
                    "error": True,
                    "message": f"File is not a DOCX document: {normalized_path}",
                    "file_path": file_path,
                }
            )

        # Read all paragraphs from DOCX
        doc = Document(normalized_path)

        full_text = ""
        total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text += paragraph.text.strip() + "\n\n"

        if not full_text.strip():
            return json.dumps(
                {
                    "content": "",
                    "message": "DOCX appears to be empty",
                    "file_path": normalized_path,
                    "total_paragraphs": total_paragraphs,
                }
            )

        # If RAG enabled, ingest into session store instead of summarizing everything
        if config.get("ENABLE_RAG", False) and _rag_available:
            tokens = count_tokens(full_text)
            if tokens > config.get("RAG", {}).get("MAX_TOKENS", 1024 * 8):
                try:
                    rag = get_rag_session()
                    if rag is None:
                        logger.warning(
                            "RAG session not initialized - falling back to summarization"
                        )
                    else:
                        num_chunks = rag.ingest(
                            os.path.basename(normalized_path),
                            full_text,
                            chunk_size_tokens=int(
                                config.get("RAG", {}).get("CHUNK_TOKENS", 1024)
                            ),
                        )

                        return json.dumps(
                            {
                                "content": "",
                                "message": "Document ingested into RAG store",
                                "file_path": normalized_path,
                                "total_paragraphs": total_paragraphs,
                                "chunks_indexed": num_chunks,
                            }
                        )
                except Exception as e:
                    logger.exception("RAG ingestion failed: %s", e)
                    # fallback to summarization

        # Process with chunking if needed (fallback or RAG not enabled)
        processed_content, metadata = await process_large_content(full_text)

        return json.dumps(
            {
                "content": processed_content.strip(),
                "file_path": normalized_path,
                "total_paragraphs": total_paragraphs,
                "processing_metadata": metadata,
            }
        )

    except Exception as e:
        logger.error(f"Error during DOCX read: {str(e)}", exc_info=True)

        return json.dumps(
            {
                "error": True,
                "message": f"Error reading DOCX: {str(e)}",
                "file_path": file_path,
            }
        )
